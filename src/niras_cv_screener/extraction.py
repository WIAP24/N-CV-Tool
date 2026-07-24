from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


@dataclass
class ExtractionResult:
    file_name: str
    sha256: str
    text: str
    blocks: List[Dict[str, Any]] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    used_ocr: bool = False

    @property
    def char_count(self) -> int:
        return len(self.text or "")


def safe_stem(file_name: str) -> str:
    stem = Path(file_name).stem
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", stem).strip("._") or "cv"


def file_sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def is_supported_file(path: Path) -> bool:
    return path.suffix.lower() in SUPPORTED_EXTENSIONS and path.is_file()


def list_cv_paths(folder: Path, recursive: bool = False) -> List[Path]:
    if not folder.is_dir():
        return []
    iterator = folder.rglob("*") if recursive else folder.iterdir()
    return sorted([path for path in iterator if is_supported_file(path)], key=lambda p: p.name.lower())


def extract_from_path(path: Path, use_ocr: bool = False) -> ExtractionResult:
    data = path.read_bytes()
    return extract_from_bytes(path.name, data, use_ocr=use_ocr)


def extract_from_bytes(file_name: str, data: bytes, use_ocr: bool = False) -> ExtractionResult:
    ext = Path(file_name).suffix.lower()
    digest = file_sha256(data)
    if ext == ".pdf":
        return extract_pdf(file_name, data, digest, use_ocr=use_ocr)
    if ext == ".docx":
        return extract_docx(file_name, data, digest)
    if ext in {".txt", ".md"}:
        text = data.decode("utf-8", errors="ignore").strip()
        warnings = [] if len(text) >= 100 else ["Text file contains very little text."]
        return ExtractionResult(file_name=file_name, sha256=digest, text=text, blocks=[{"source_ref": "text", "text": text}], warnings=warnings)
    raise ValueError(f"Unsupported file type: {ext}")


def extract_pdf(file_name: str, data: bytes, digest: str, use_ocr: bool = False) -> ExtractionResult:
    import io

    from pypdf import PdfReader

    warnings: List[str] = []
    blocks: List[Dict[str, Any]] = []
    reader = PdfReader(io.BytesIO(data))
    if reader.is_encrypted:
        try:
            reader.decrypt("")
            warnings.append("PDF is encrypted; attempted empty-password decrypt.")
        except Exception as exc:
            raise ValueError(f"Could not read encrypted PDF: {exc}") from exc

    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = clean_text(text)
        if text:
            blocks.append({"source_ref": f"page {page_number}", "page": page_number, "text": text})

    used_ocr = False
    if use_ocr and not enough_text(blocks):
        ocr_blocks, ocr_warnings = ocr_pdf(data)
        warnings.extend(ocr_warnings)
        if enough_text(ocr_blocks):
            blocks = ocr_blocks
            used_ocr = True

    combined = "\n\n".join(f"[{block['source_ref']}]\n{block['text']}" for block in blocks).strip()
    if not combined:
        warnings.append("No extractable text found. This is probably a scanned/image-only PDF and needs OCR.")
    elif len(combined) < 500:
        warnings.append("Very little text extracted. Check whether the CV is scanned or image-heavy.")
    return ExtractionResult(file_name=file_name, sha256=digest, text=combined, blocks=blocks, warnings=dedupe(warnings), used_ocr=used_ocr)


def ocr_pdf(data: bytes) -> tuple[List[Dict[str, Any]], List[str]]:
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
    except Exception:
        return [], ["OCR was requested, but optional OCR packages are not installed. Install requirements-ocr.txt and Tesseract OCR."]

    warnings: List[str] = []
    blocks: List[Dict[str, Any]] = []
    try:
        pages = convert_from_bytes(data, dpi=200)
    except Exception as exc:
        return [], [f"OCR could not render PDF pages: {exc}"]

    for page_number, image in enumerate(pages, start=1):
        try:
            text = clean_text(pytesseract.image_to_string(image))
            if text:
                blocks.append({"source_ref": f"ocr page {page_number}", "page": page_number, "text": text})
        except Exception as exc:
            warnings.append(f"OCR failed on page {page_number}: {exc}")
    if not blocks:
        warnings.append("OCR did not return usable text.")
    return blocks, warnings


def extract_docx(file_name: str, data: bytes, digest: str) -> ExtractionResult:
    import io

    from docx import Document

    warnings: List[str] = []
    blocks: List[Dict[str, Any]] = []
    doc = Document(io.BytesIO(data))

    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = clean_text(paragraph.text)
        if text:
            blocks.append({"source_ref": f"paragraph {index}", "text": text})

    for table_index, table in enumerate(doc.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            cells = [clean_text(cell.text) for cell in row.cells]
            text = " | ".join(cell for cell in cells if cell)
            if text:
                blocks.append({"source_ref": f"table {table_index} row {row_index}", "text": text})

    combined = "\n".join(f"[{block['source_ref']}] {block['text']}" for block in blocks).strip()
    if not combined:
        warnings.append("No text found in DOCX.")
    elif len(combined) < 500:
        warnings.append("Very little text extracted from DOCX.")
    return ExtractionResult(file_name=file_name, sha256=digest, text=combined, blocks=blocks, warnings=warnings)


def enough_text(blocks: List[Dict[str, Any]]) -> bool:
    return sum(len(str(block.get("text", ""))) for block in blocks) >= 500


def dedupe(values: List[str]) -> List[str]:
    out: List[str] = []
    seen: set[str] = set()
    for value in values:
        if value not in seen:
            out.append(value)
            seen.add(value)
    return out


def clean_text(value: str) -> str:
    value = value or ""
    value = value.encode("utf-8", errors="ignore").decode("utf-8")
    value = value.replace("\x00", " ")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()
